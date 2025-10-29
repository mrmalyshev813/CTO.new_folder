import logging
import os
import uuid
from pathlib import Path
from typing import Tuple, Optional
from docx import Document
from docx.shared import Pt
from weasyprint import HTML

logger = logging.getLogger(__name__)

TEMP_DIR = Path("/tmp/adlook_exports")
TEMP_DIR.mkdir(exist_ok=True)


def create_docx(proposal_text: str, analysis_id: str) -> Tuple[Optional[str], bool, Optional[str]]:
    """
    Create a DOCX file from proposal text.
    
    Args:
        proposal_text: The proposal text to convert
        analysis_id: Unique identifier for this analysis
        
    Returns:
        Tuple of (file_path, success, error_message)
    """
    try:
        doc = Document()
        
        lines = proposal_text.split('\n')
        
        for line in lines:
            if line.strip():
                paragraph = doc.add_paragraph(line)
                for run in paragraph.runs:
                    run.font.size = Pt(11)
            else:
                doc.add_paragraph()
        
        file_path = TEMP_DIR / f"{analysis_id}.docx"
        doc.save(str(file_path))
        
        logger.info(f"Created DOCX file: {file_path}")
        return str(file_path), True, None
        
    except Exception as e:
        error_msg = f"Error creating DOCX: {str(e)}"
        logger.error(error_msg)
        return None, False, error_msg


def create_pdf(proposal_text: str, analysis_id: str) -> Tuple[Optional[str], bool, Optional[str]]:
    """
    Create a PDF file from proposal text.
    
    Args:
        proposal_text: The proposal text to convert
        analysis_id: Unique identifier for this analysis
        
    Returns:
        Tuple of (file_path, success, error_message)
    """
    try:
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <style>
                body {{
                    font-family: Arial, sans-serif;
                    font-size: 11pt;
                    line-height: 1.6;
                    margin: 40px;
                    color: #333;
                }}
                p {{
                    margin: 8px 0;
                }}
                .empty-line {{
                    margin: 4px 0;
                }}
            </style>
        </head>
        <body>
        """
        
        lines = proposal_text.split('\n')
        for line in lines:
            if line.strip():
                escaped_line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                html_content += f"<p>{escaped_line}</p>\n"
            else:
                html_content += '<p class="empty-line">&nbsp;</p>\n'
        
        html_content += """
        </body>
        </html>
        """
        
        file_path = TEMP_DIR / f"{analysis_id}.pdf"
        HTML(string=html_content).write_pdf(str(file_path))
        
        logger.info(f"Created PDF file: {file_path}")
        return str(file_path), True, None
        
    except Exception as e:
        error_msg = f"Error creating PDF: {str(e)}"
        logger.error(error_msg)
        return None, False, error_msg


def get_file_path(analysis_id: str, file_type: str) -> Optional[str]:
    """
    Get the file path for a given analysis ID and file type.
    
    Args:
        analysis_id: The analysis identifier
        file_type: Either 'docx' or 'pdf'
        
    Returns:
        File path if exists, None otherwise
    """
    file_path = TEMP_DIR / f"{analysis_id}.{file_type}"
    if file_path.exists():
        return str(file_path)
    return None
